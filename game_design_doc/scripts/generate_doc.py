#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
游戏功能设计文档生成器
根据模板自动生成标准化的Word文档框架
"""

import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
import os
from datetime import datetime


class GameDocGenerator:
    """游戏功能文档生成器"""
    
    def __init__(self, func_name, func_type, output_path=None):
        self.func_name = func_name
        self.func_type = func_type
        self.output_path = output_path or f"{func_name}_设计文档.docx"
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置中文字体
        self.doc.styles['Normal'].font.name = '宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        self.doc.styles['Normal'].font.size = Pt(12)
        
    def _add_heading(self, text, level=1):
        """添加标题"""
        heading = self.doc.add_heading(text, level=level)
        heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        return heading
    
    def _add_paragraph(self, text, style=None):
        """添加段落"""
        para = self.doc.add_paragraph(text, style=style)
        return para
    
    def _add_table(self, headers, rows_data=None, num_empty_rows=3):
        """添加表格"""
        num_rows = len(rows_data) + 1 if rows_data else num_empty_rows + 1
        num_cols = len(headers)
        
        table = self.doc.add_table(rows=num_rows, cols=num_cols)
        table.style = 'Table Grid'
        
        # 设置表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # 表头加粗
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # 填充数据
        if rows_data:
            for i, row_data in enumerate(rows_data, start=1):
                row_cells = table.rows[i].cells
                for j, cell_data in enumerate(row_data):
                    row_cells[j].text = str(cell_data)
        
        return table
    
    def generate(self):
        """生成文档"""
        # 标题
        title = self.doc.add_heading(f'{self.func_name} 功能设计文档', level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加生成信息
        self._add_paragraph(f'生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
        self._add_paragraph(f'功能类型：{self._get_type_name()}')
        self._add_paragraph('')
        
        # 一、设计目的
        self._add_section_design_purpose()
        
        # 二、功能概述
        self._add_section_overview()
        
        # 三、规则说明
        self._add_section_rules()
        
        # 四、策划需求
        self._add_section_requirements()
        
        # 保存文档
        self.doc.save(self.output_path)
        print(f"✅ 文档已生成：{self.output_path}")
    
    def _get_type_name(self):
        """获取功能类型中文名"""
        type_map = {
            'system': '系统玩法',
            'building': '建筑功能',
            'activity': '活动功能',
            'other': '其他'
        }
        return type_map.get(self.func_type, '未知')
    
    def _add_section_design_purpose(self):
        """添加：一、设计目的"""
        self._add_heading('一、设计目的', level=1)
        
        self._add_heading('1.1 功能定位', level=2)
        self._add_paragraph('[说明功能在游戏中的定位和作用，解决的核心问题]')
        self._add_paragraph('')
        
        self._add_heading('1.2 期望体验', level=2)
        self._add_paragraph('[描述玩家使用该功能时的预期体验和价值]')
        self._add_paragraph('')
    
    def _add_section_overview(self):
        """添加：二、功能概述"""
        self._add_heading('二、功能概述', level=1)
        
        self._add_heading('2.1 背景概述', level=2)
        self._add_paragraph('[功能的背景故事或世界观设定]')
        self._add_paragraph('')
        
        self._add_heading('2.2 功能简介', level=2)
        self._add_paragraph('核心玩法：[用1-3段话描述功能的核心玩法流程]')
        self._add_paragraph('')
        self._add_paragraph('主要特点：')
        for i in range(1, 4):
            self._add_paragraph(f'{i}. [关键特性{i}]', style='List Number')
        self._add_paragraph('')
        
        self._add_heading('2.3 结构划分', level=2)
        self._add_paragraph('[使用列表或文字描述功能的结构组成]')
        self._add_paragraph('示例：')
        self._add_paragraph('功能名称', style='List Bullet')
        self._add_paragraph('模块A', style='List Bullet 2')
        self._add_paragraph('模块B', style='List Bullet 2')
        self._add_paragraph('')
    
    def _add_section_rules(self):
        """添加：三、规则说明"""
        self._add_heading('三、规则说明', level=1)
        
        # 添加重要提示
        self._add_paragraph('⚠️ 重要原则：禁止使用代码和伪代码')
        self._add_paragraph('在描述客户端和服务器的规则时，严禁使用任何形式的代码或伪代码。必须使用纯文本、表格、列表来描述规则。')
        self._add_paragraph('')
        self._add_paragraph('⚠️ 术语标注要求：所有核心功能、游戏元素、操作术语必须使用【】符号标注，保持与项目已有术语一致。')
        self._add_paragraph('')
        
        # 根据功能类型添加不同的规则章节
        if self.func_type == 'activity':
            self._add_activity_rules()
        elif self.func_type == 'building':
            self._add_building_rules()
        else:
            self._add_common_rules()
    
    def _add_common_rules(self):
        """添加通用规则"""
        self._add_heading('3.1 开启条件', level=2)
        self._add_paragraph('[列出功能解锁的所有条件]')
        self._add_table(['条件类型', '具体要求', '说明'])
        self._add_paragraph('')
        
        self._add_heading('3.2 参与条件', level=2)
        self._add_paragraph('[描述玩家进入功能或参与玩法的条件]')
        self._add_paragraph('')
        
        self._add_heading('3.3 运行规则', level=2)
        self._add_paragraph('[描述功能的核心运行逻辑]')
        self._add_paragraph('')
        
        self._add_heading('3.4 特殊处理', level=2)
        self._add_paragraph('[列出所有特殊情况及其处理方式]')
        self._add_table(['特殊情况', '处理方式'])
        self._add_paragraph('')
    
    def _add_activity_rules(self):
        """添加活动类功能规则"""
        self._add_heading('3.1 活动状态机', level=2)
        self._add_table(['状态', '说明', '进入条件', '退出条件'])
        self._add_paragraph('')
        
        self._add_heading('3.2 开启条件', level=2)
        self._add_table(['条件类型', '具体要求', '说明'])
        self._add_paragraph('')
        
        self._add_heading('3.3 参与条件', level=2)
        self._add_paragraph('[描述玩家参与活动的条件]')
        self._add_paragraph('')
        
        self._add_heading('3.4 循环方式', level=2)
        self._add_table(['循环方式', '循环规则', '案例'])
        self._add_paragraph('')
        
        self._add_heading('3.5 结束规则', level=2)
        self._add_table(['结束条件', '结束规则', '后续处理'])
        self._add_paragraph('')
        
        self._add_heading('3.6 特殊处理', level=2)
        self._add_table(['特殊情况', '处理方式'])
        self._add_paragraph('')
        
        self._add_heading('3.7 红点提示规则', level=2)
        self._add_table(['提示位置', '出现条件', '消失条件'])
        self._add_paragraph('')
    
    def _add_building_rules(self):
        """添加建筑类功能规则"""
        self._add_heading('3.1 建筑初始状态', level=2)
        self._add_paragraph('[描述建筑的初始状态和默认配置]')
        self._add_paragraph('')
        
        self._add_heading('3.2 建筑解锁条件', level=2)
        self._add_table(['建筑名称', '解锁条件', '说明'])
        self._add_paragraph('')
        
        self._add_heading('3.3 升级规则', level=2)
        self._add_paragraph('[描述建筑升级的条件和流程]')
        self._add_table(['升级条件', '说明'])
        self._add_paragraph('')
        
        self._add_heading('3.4 加速规则', level=2)
        self._add_paragraph('[描述加速道具使用规则和钻石加速规则]')
        self._add_paragraph('')
        
        self._add_heading('3.5 建造/升级表现', level=2)
        self._add_paragraph('[描述建筑建造和升级时的客户端表现]')
        self._add_paragraph('')
        
        self._add_heading('3.6 特殊处理', level=2)
        self._add_table(['特殊情况', '处理方式'])
        self._add_paragraph('')
    
    def _add_section_requirements(self):
        """添加：四、策划需求"""
        self._add_heading('四、策划需求', level=1)
        
        self._add_heading('4.1 数值需求', level=2)
        self._add_paragraph('⚠️ 必须明确区分硬编码参数和可配置参数')
        self._add_paragraph('')
        self._add_paragraph('硬编码参数（固定值，不需要在配置表中存储）：')
        self._add_paragraph('- [参数名称]：[固定值]', style='List Bullet')
        self._add_paragraph('')
        self._add_paragraph('可配置参数（需要在配置表中设计字段）：')
        self._add_table(['参数名称', '取值', '说明', '配置表字段'])
        self._add_paragraph('')
        
        self._add_heading('4.2 系统需求', level=2)
        self._add_paragraph('[说明需要其他系统提供的支持，使用【】标注系统名称]')
        self._add_paragraph('示例：')
        self._add_paragraph('需要【背包系统】支持道具存储和使用', style='List Bullet')
        self._add_paragraph('需要【任务系统】提供任务进度追踪接口', style='List Bullet')
        self._add_paragraph('')
        
        self._add_heading('4.3 配置表需求', level=2)
        self._add_paragraph('⚠️ 配置表复用原则：优先复用已有配置表，禁止重复创建。只有在没有合适的已有表时，才能创建新表。')
        self._add_paragraph('')
        self._add_paragraph('【增加数据】在已有表 xxx_config 中增加以下数据行：')
        self._add_table(['字段名', '数据值示例'], num_empty_rows=2)
        self._add_paragraph('')
        self._add_paragraph('【增加字段】在已有表 xxx_config 中增加以下字段：')
        self._add_table(['字段名', '类型', '说明', '对应规则参数'], num_empty_rows=2)
        self._add_paragraph('')
        self._add_paragraph('【新建】新建表：xxx_config')
        self._add_paragraph('说明为什么需要新建：[现有的 xxx 表都无法满足该功能的配置需求，因为...]')
        self._add_table(['字段名', '类型', '说明', '对应规则参数'])
        self._add_paragraph('')


def main():
    """主函数"""
    parser = argparse.ArgumentParser(description='游戏功能设计文档生成器')
    parser.add_argument('--name', required=True, help='功能名称')
    parser.add_argument('--type', required=True, 
                        choices=['system', 'building', 'activity', 'other'],
                        help='功能类型：system(系统玩法)/building(建筑)/activity(活动)/other(其他)')
    parser.add_argument('--output', help='输出文件路径（可选）')
    
    args = parser.parse_args()
    
    # 创建生成器并生成文档
    generator = GameDocGenerator(args.name, args.type, args.output)
    generator.generate()
    
    print(f"\n📄 文档生成完成！")
    print(f"📁 文件位置：{os.path.abspath(generator.output_path)}")
    print(f"\n💡 提示：请打开文档并根据实际需求填充各章节内容。")


if __name__ == '__main__':
    main()
