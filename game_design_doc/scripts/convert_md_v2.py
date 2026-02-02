
import os
import re
import argparse
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

class MarkdownToDocx:
    def __init__(self, input_file, output_file):
        self.input_file = input_file
        self.output_file = output_file
        self.doc = Document()
        self._setup_styles()
        
    def _setup_styles(self):
        """配置文档基本样式，支持中文"""
        style = self.doc.styles['Normal']
        style.font.name = '宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        style.font.size = Pt(12)
        
        # 设置标题样式
        for i in range(1, 10):
            if f'Heading {i}' in self.doc.styles:
                h_style = self.doc.styles[f'Heading {i}']
                h_style.font.name = '微软雅黑'
                h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                h_style.font.color.rgb = RGBColor(0, 0, 0) # 黑色
                h_style.font.bold = True


    def parse_inline_styles(self, paragraph, text):
        """解析行内样式：**Bold**, *Italic*"""
        # 使用更复杂的正则同时匹配 **bold** 和 *italic*
        # 注意：这里简单的实现不支持嵌套
        parts = re.split(r'(\*\*.*?\*\*|\*[^*]+?\*)', text)
        
        for part in parts:
            if not part:
                continue
                
            if part.startswith('**') and part.endswith('**'):
                # 加粗 (Bold)
                content = part[2:-2]
                if content:
                    run = paragraph.add_run(content)
                    run.font.bold = True
                    run.font.name = '宋体'
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            elif part.startswith('*') and part.endswith('*'):
                # 斜体 (Italic) -> 映射为楷体 (KaiTi)
                content = part[1:-1]
                if content:
                    run = paragraph.add_run(content)
                    # run.font.italic = True # 不使用斜体，而是使用楷体
                    run.font.name = '楷体' 
                    run.element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            else:
                # 普通文本
                run = paragraph.add_run(part)
                run.font.name = '宋体'
                run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


    def convert(self):
        with open(self.input_file, 'r', encoding='utf-8') as f:
            lines = f.readlines()

        table_mode = False
        table_rows = []
        code_mode = False

        for line in lines:
            line = line.strip()
            
            # 处理代码块
            if line.startswith('```'):
                code_mode = not code_mode
                continue
            
            if code_mode:
                p = self.doc.add_paragraph()
                p.style = 'No Spacing'
                run = p.add_run(line)
                run.font.name = 'Courier New'
                continue

            # 处理表格
            if line.startswith('|') and line.endswith('|'):
                if not table_mode:
                    table_mode = True
                    table_rows = []
                
                # 跳过分隔符行 |---|---|
                if '---' in line:
                    continue
                    
                row_data = [cell.strip() for cell in line.strip('|').split('|')]
                table_rows.append(row_data)
                continue
            else:
                if table_mode:
                    # 表格结束，写入表格
                    self._create_table(table_rows)
                    table_mode = False
                    table_rows = []

            if not line:
                continue

            # 处理标题
            header_match = re.match(r'^(#{1,6})\s+(.*)', line)
            if header_match:
                level = len(header_match.group(1))
                content = header_match.group(2)
                self.doc.add_heading(content, level=level)
                continue

            # 处理列表
            if line.startswith('- ') or line.startswith('* '):
                p = self.doc.add_paragraph(style='List Bullet')
                self.parse_inline_styles(p, line[2:])
                continue
            
            # 处理有序列表 (简单匹配 1. )
            if re.match(r'^\d+\.\s', line):
                content = re.sub(r'^\d+\.\s', '', line)
                p = self.doc.add_paragraph(style='List Number')
                self.parse_inline_styles(p, content)
                continue

            # 普通段落
            p = self.doc.add_paragraph()
            self.parse_inline_styles(p, line)

        # 如果文件以表格结尾
        if table_mode:
             self._create_table(table_rows)

        self.doc.save(self.output_file)
        print(f"✅ Converted: {self.output_file}")

    def _create_table(self, rows):
        if not rows:
            return
        
        num_cols = max(len(r) for r in rows)
        table = self.doc.add_table(rows=len(rows), cols=num_cols)
        table.style = 'Table Grid'
        
        for i, row_data in enumerate(rows):
            row_cells = table.rows[i].cells
            for j, cell_text in enumerate(row_data):
                if j < len(row_cells):
                    cell = row_cells[j]
                    # 清除可能默认存在的段落
                    # cell.text = "" 
                    # 应该保留第一个段落但清空内容，或者直接获取第一个段落
                    if not cell.paragraphs:
                        p = cell.add_paragraph()
                    else:
                        p = cell.paragraphs[0]
                        p.clear() # 清空内容
                        
                    self.parse_inline_styles(p, cell_text)
                    
                    # 第一行 (表头) 样式优化
                    if i == 0:
                        # 背景色 (灰色)
                        tcPr = cell._element.get_or_add_tcPr()
                        shd = OxmlElement('w:shd')
                        shd.set(qn('w:val'), 'clear')
                        shd.set(qn('w:color'), 'auto')
                        shd.set(qn('w:fill'), 'F2F2F2') # 浅灰色背景
                        tcPr.append(shd)
                        
                        # 字体加粗 + 微软雅黑
                        for run in p.runs:
                            run.font.bold = True
                            run.font.name = '微软雅黑'
                            run.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def main():
    parser = argparse.ArgumentParser(description='Convert Markdown to Docx (Custom)')
    parser.add_argument('input', help='Input Markdown file')
    parser.add_argument('output', help='Output Docx file')
    args = parser.parse_args()
    
    converter = MarkdownToDocx(args.input, args.output)
    converter.convert()

if __name__ == "__main__":
    main()
