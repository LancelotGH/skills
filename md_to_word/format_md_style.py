import os
import sys
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

input_docx = sys.argv[1]
output_docx = sys.argv[2]

try:
    doc = Document(input_docx)
except Exception as e:
    print(f"Failed to open document: {e}")
    sys.exit(1)

def set_font(style, ascii_font, eastasia_font, size_pt, bold=None, color_rgb=None):
    if style is None: return
    font = style.font
    font.name = ascii_font
    font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold
    if color_rgb is not None:
        font.color.rgb = color_rgb
    
    if style.element.rPr is not None:
        rFonts = style.element.rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = style.element.rPr.makeelement(qn('w:rFonts'))
            style.element.rPr.append(rFonts)
        rFonts.set(qn('w:eastAsia'), eastasia_font)

# 1. Update general styles
try:
    # Use standard headings but keep compact spacing
    set_font(doc.styles['Heading 1'], 'Microsoft YaHei', 'Microsoft YaHei', 18, True, RGBColor(0, 0, 0))
    set_font(doc.styles['Heading 2'], 'Microsoft YaHei', 'Microsoft YaHei', 15, True, RGBColor(0, 0, 0))
    set_font(doc.styles['Heading 3'], 'Microsoft YaHei', 'Microsoft YaHei', 12, True, RGBColor(0, 0, 0))
    set_font(doc.styles['Heading 4'], 'Microsoft YaHei', 'Microsoft YaHei', 11, True, RGBColor(0, 0, 0))
    set_font(doc.styles['Normal'], 'Microsoft YaHei', 'Microsoft YaHei', 10, False, RGBColor(40, 40, 40))
    
    doc.styles['Normal'].paragraph_format.line_spacing = 1.15
    doc.styles['Normal'].paragraph_format.space_after = Pt(6)

    # Disable default italics on Block Text (quotations) often set by pandoc
    if 'Block Text' in doc.styles:
         doc.styles['Block Text'].font.italic = False
         
except Exception as e:
    print(f"Style warning: {e}")

# Apply specific styling for mimicking MD elements
def style_as_quote(paragraph):
    # Left border and indent for quotes
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr')
        pPr.append(pbdr)
    left_bdr = OxmlElement('w:left')
    left_bdr.set(qn('w:val'), 'single')
    left_bdr.set(qn('w:sz'), '12') # thickness
    left_bdr.set(qn('w:space'), '4')
    left_bdr.set(qn('w:color'), 'CCCCCC')
    pbdr.append(left_bdr)
    
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(4)
    
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor(90, 90, 90)

def style_as_code_block(paragraph):
    # Grey background mimicking code block
    pPr = paragraph._p.get_or_add_pPr()
    shd = pPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        pPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F6F8FA')

    pbdr = pPr.find(qn('w:pBdr'))
    if pbdr is None:
        pbdr = OxmlElement('w:pBdr')
        pPr.append(pbdr)
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') 
        border.set(qn('w:space'), '1')
        border.set(qn('w:color'), 'E1E4E8')
        pbdr.append(border)

    paragraph.paragraph_format.left_indent = Inches(0.1)
    paragraph.paragraph_format.right_indent = Inches(0.1)
    
    for run in paragraph.runs:
        run.font.name = 'Consolas' # Code font
        run.font.size = Pt(9.5)

def style_as_inline_code(run):
    run.font.name = 'Consolas'
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(199, 37, 78)
    
    rPr = run.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rPr.makeelement(qn('w:rFonts'))
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

    shd = rPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        rPr.append(shd)
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F9F2F4')

# 2. Iterate paragraphs to wipe all italics and set specific MD styles
for i, paragraph in enumerate(doc.paragraphs):
    if paragraph.style.name in ['List Paragraph', 'Compact']:
        paragraph.paragraph_format.space_before = Pt(2)
        paragraph.paragraph_format.space_after = Pt(2)
        paragraph.paragraph_format.line_spacing = 1.15

    # Remove any stray italic settings from runs
    for run in paragraph.runs:
        run.italic = False
        
        # Keep Microsoft YaHei as default text font (unless it's code block)
        if paragraph.style.name != 'Source Code':
            run.font.name = 'Microsoft YaHei'
            rPr = run.element.get_or_add_rPr()
            rFonts = rPr.find(qn('w:rFonts'))
            if rFonts is None:
                rFonts = rPr.makeelement(qn('w:rFonts'))
                rPr.append(rFonts)
            rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    # Identify typical Pandoc styles for MD elements
    if paragraph.style.name.startswith('Heading'):
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(4)
    elif paragraph.style.name in ['Block Text', 'Quote']:
        style_as_quote(paragraph)
    elif paragraph.style.name == 'Source Code' or 'Code' in paragraph.style.name:
        style_as_code_block(paragraph)
    
    # Process inline code
    for run in paragraph.runs:
        if run.style and run.style.name == 'Verbatim Char':
            style_as_inline_code(run)

# 3. Optimize Tables for better appearance
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_cell_background(cell, fill_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill_color)

for table in doc.tables:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            # Set fine gray borders
            set_cell_border(
                cell,
                top={"sz": 4, "val": "single", "color": "BFBFBF"},
                bottom={"sz": 4, "val": "single", "color": "BFBFBF"},
                start={"sz": 4, "val": "single", "color": "BFBFBF"},
                end={"sz": 4, "val": "single", "color": "BFBFBF"},
            )
            # Shade header row using a light blue
            if row_idx == 0:
                set_cell_background(cell, "E6F0FA")
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.font.bold = True
                        run.italic = False # Ensure no italic in header
            else:
                for p in cell.paragraphs:
                    # Alternating colors: alternate rows slightly gray
                    if row_idx % 2 == 1:
                        set_cell_background(cell, "F9F9F9")
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Make compact table contents
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.line_spacing = 1.0
                for run in p.runs:
                    run.italic = False # Clean italic
                    run.font.name = 'Microsoft YaHei'
                    run.font.size = Pt(9.5) # slightly smaller font inside tables
                    rPr = run.element.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = rPr.makeelement(qn('w:rFonts'))
                        rPr.append(rFonts)
                    rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

# 4. Final Global Style Sweep: Eradicate all Italics
def eradicate_italics(document):
    try:
        for style in document.styles:
            if hasattr(style, 'font') and style.font is not None:
                style.font.italic = False
    except Exception as e:
        print(f"Warn: {e}")

    for p in document.paragraphs:
        for r in p.runs:
            r.italic = False
            r.font.italic = False
            
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.italic = False
                        r.font.italic = False

eradicate_italics(doc)

doc.save(output_docx)
print(f"Optimization complete for {output_docx}.")
